from __future__ import annotations

from dataclasses import dataclass
from itertools import count
from typing import Any
from uuid import uuid4

from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

from app import models
from app.core.constants import SYSTEM_FREE_WRITE_STYLE_ID
from app.core.generation_policy import GenerationMode
from app.core.model_gateway import ModelGateway, ModelResult
from app.core.prompt_composer import WritingTaskInput, compose_prompt, compose_free_prompt
from app.core.style_profile_builder import build_style_profile_v2
from app.core.text_parser import split_paragraphs


DEMO_USER_ID = "demo_user"
FALLBACK_GENERATION_COUNTER = count(1)


def new_id(prefix: str) -> str:
    return f"{prefix}_{uuid4().hex[:12]}"


def ensure_demo_user(db: Session) -> models.User:
    user = db.get(models.User, DEMO_USER_ID)
    if user:
        return user
    user = models.User(id=DEMO_USER_ID, display_name="Demo 用户", mode="demo")
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


def _require_user_materials(db: Session, *, user_id: str, material_ids: list[str]) -> list[models.Material]:
    if not material_ids:
        raise HTTPException(status_code=400, detail="material_ids is required")
    materials = list(
        db.scalars(
            select(models.Material)
            .where(models.Material.user_id == user_id, models.Material.id.in_(material_ids))
            .options(selectinload(models.Material.paragraphs))
        )
    )
    if len(materials) != len(set(material_ids)):
        raise HTTPException(status_code=404, detail="one or more materials were not found")
    return materials


def create_material(
    db: Session,
    *,
    user_id: str,
    title: str,
    genre: str,
    source_type: str,
    source_filename: str | None,
    content: str,
) -> models.Material:
    normalized_title = title.strip() or (source_filename or "未命名作品")
    normalized_genre = genre.strip() or "散文"
    paragraphs = split_paragraphs(content)
    if not paragraphs:
        raise HTTPException(status_code=400, detail="material content is empty")
    material = models.Material(
        id=new_id("mat"),
        user_id=user_id,
        title=normalized_title,
        genre=normalized_genre,
        source_type=source_type,
        source_filename=source_filename,
        content="\n\n".join(paragraphs),
        char_count=sum(len(item) for item in paragraphs),
        paragraph_count=len(paragraphs),
    )
    material.paragraphs = [
        models.MaterialParagraph(
            id=new_id("mpara"),
            position=index,
            content=paragraph,
            char_count=len(paragraph),
        )
        for index, paragraph in enumerate(paragraphs, start=1)
    ]
    db.add(material)
    db.commit()
    db.refresh(material)
    return material


def list_materials(db: Session, *, user_id: str) -> list[models.Material]:
    return list(
        db.scalars(
            select(models.Material)
            .where(models.Material.user_id == user_id)
            .options(selectinload(models.Material.paragraphs))
            .order_by(models.Material.created_at.desc())
        )
    )


def build_style_draft(materials: list[models.Material]) -> dict[str, Any]:
    return build_style_profile_v2(materials)


def create_style_analysis_job(db: Session, *, user_id: str, material_ids: list[str]) -> models.StyleAnalysisJob:
    materials = _require_user_materials(db, user_id=user_id, material_ids=material_ids)
    job = models.StyleAnalysisJob(
        id=new_id("style_job"),
        user_id=user_id,
        material_ids=[item.id for item in materials],
        status="draft_pending_confirmation",
        draft_profile=build_style_draft(materials),
    )
    db.add(job)
    db.commit()
    db.refresh(job)
    return job


def get_style_analysis_job(db: Session, *, user_id: str, job_id: str) -> models.StyleAnalysisJob:
    job = db.get(models.StyleAnalysisJob, job_id)
    if not job or job.user_id != user_id:
        raise HTTPException(status_code=404, detail="style analysis job not found")
    return job


def confirm_style_profile(
    db: Session,
    *,
    user_id: str,
    job_id: str,
    name: str,
    profile: dict[str, Any] | None,
) -> models.StyleProfile:
    job = get_style_analysis_job(db, user_id=user_id, job_id=job_id)
    normalized_name = name.strip()
    if not normalized_name:
        raise HTTPException(status_code=400, detail="style name is required")
    existing_for_job = db.scalar(
        select(models.StyleProfile).where(
            models.StyleProfile.user_id == user_id,
            models.StyleProfile.source_job_id == job.id,
            models.StyleProfile.status == "active",
        )
    )
    if job.status == "confirmed" and existing_for_job:
        return existing_for_job
    if job.status != "draft_pending_confirmation":
        raise HTTPException(status_code=409, detail="style analysis job is not waiting for confirmation")
    existing_style = db.scalar(
        select(models.StyleProfile).where(
            models.StyleProfile.user_id == user_id,
            models.StyleProfile.status == "active",
            models.StyleProfile.name == normalized_name,
        )
    )
    if existing_style:
        raise HTTPException(status_code=409, detail="风格名称已存在，请换一个名称。")
    final_profile = profile or job.draft_profile
    final_profile["style_name"] = normalized_name
    final_profile["status"] = "active"
    style = models.StyleProfile(
        id=new_id("style"),
        user_id=user_id,
        source_job_id=job.id,
        name=normalized_name,
        status="active",
        profile=final_profile,
    )
    job.status = "confirmed"
    db.add(style)
    # 风格生成完成后推送系统通知（融合进消息中心）
    from app.core import message_service as _ms

    _ms._add_system_message(
        db,
        user_id,
        title="风格生成完成",
        body=f"您上传的作品已分析完成，风格「{normalized_name}」已生成，现在可以开始用该风格写作了。",
        category="system",
    )
    db.commit()
    db.refresh(style)
    return style


def list_style_profiles(db: Session, *, user_id: str | None) -> list[models.StyleProfile]:
    conditions = [models.StyleProfile.status == "active"]
    if user_id is not None:
        conditions.append(
            (models.StyleProfile.user_id == user_id) | (models.StyleProfile.is_recommended == True)
        )
    else:
        conditions.append(models.StyleProfile.is_recommended == True)
    return list(
        db.scalars(
            select(models.StyleProfile)
            .where(*conditions)
            .order_by(models.StyleProfile.is_recommended.desc(), models.StyleProfile.created_at.desc())
        )
    )


def delete_style_profile(db: Session, *, user_id: str, style_profile_id: str) -> models.StyleProfile:
    style = db.get(models.StyleProfile, style_profile_id)
    if not style or style.user_id != user_id:
        raise HTTPException(status_code=404, detail="style profile not found")
    if style.status != "deleted":
        style.status = "deleted"
        if isinstance(style.profile, dict):
            style.profile = {**style.profile, "status": "deleted"}
        db.commit()
        db.refresh(style)
    return style


def get_active_style(db: Session, *, user_id: str, style_profile_id: str) -> models.StyleProfile:
    style = db.get(models.StyleProfile, style_profile_id)
    # 用户可使用自己的 active 风格，也可使用系统推荐风格。
    if not style or (style.user_id != user_id and not style.is_recommended):
        raise HTTPException(status_code=404, detail="style profile not found")
    if style.status != "active":
        raise HTTPException(status_code=409, detail="style profile is not active")
    return style


def update_style_profile(
    db: Session,
    *,
    user_id: str,
    style_profile_id: str,
    name: str,
    description: str | None = None,
    profile: dict[str, Any] | None,
    is_recommended: bool | None = None,
    is_admin: bool = False,
) -> models.StyleProfile:
    style = get_active_style(db, user_id=user_id, style_profile_id=style_profile_id)
    normalized_name = name.strip()
    if not normalized_name:
        raise HTTPException(status_code=400, detail="style name is required")
    duplicate = db.scalar(
        select(models.StyleProfile).where(
            models.StyleProfile.user_id == user_id,
            models.StyleProfile.status == "active",
            models.StyleProfile.name == normalized_name,
            models.StyleProfile.id != style_profile_id,
        )
    )
    if duplicate:
        raise HTTPException(status_code=409, detail="风格名称已存在，请换一个名称。")
    style.name = normalized_name
    style.description = description.strip() if description else None
    if profile is not None:
        style.profile = profile
        if isinstance(profile, dict):
            profile["style_name"] = normalized_name
    if is_recommended is not None:
        if not is_admin:
            raise HTTPException(status_code=403, detail="只有管理员可以设置推荐风格")
        style.is_recommended = is_recommended
    style.updated_at = models.utc_now()
    db.commit()
    db.refresh(style)
    return style


def set_default_style_profile(
    db: Session,
    *,
    user_id: str,
    style_profile_id: str,
) -> models.StyleProfile:
    style = get_active_style(db, user_id=user_id, style_profile_id=style_profile_id)
    if style.is_recommended:
        raise HTTPException(status_code=400, detail="推荐风格不能设为默认")
    if style.is_default:
        return style
    db.query(models.StyleProfile).filter(
        models.StyleProfile.user_id == user_id,
        models.StyleProfile.id != style_profile_id,
    ).update({models.StyleProfile.is_default: False})
    style.is_default = True
    style.updated_at = models.utc_now()
    db.commit()
    db.refresh(style)
    return style


@dataclass(frozen=True)
class CreatedWriting:
    task: models.WritingTask
    document: models.Document
    model_result: ModelResult


def create_writing(
    db: Session,
    *,
    user_id: str,
    style_profile_id: str,
    task_input: WritingTaskInput,
    requested_mode: GenerationMode | str | None = None,
    rag_snippets: list[str] | None = None,
) -> CreatedWriting:
    style = get_active_style(db, user_id=user_id, style_profile_id=style_profile_id)
    prompt = compose_prompt(
        task=task_input,
        style_profile=style.profile,
        requested_mode=requested_mode or GenerationMode.STYLE_PROMPT_ONLY,
        rag_snippets=rag_snippets,
    )
    fallback = _fallback_article(style_name=style.name, task=task_input)
    model_result = ModelGateway().generate(
        messages=[
            {"role": "system", "content": prompt.system_prompt},
            {"role": "user", "content": prompt.user_prompt},
        ],
        purpose="writing",
        fallback=fallback,
    )
    paragraphs = split_paragraphs(model_result.content)
    document = models.Document(
        id=new_id("doc"),
        user_id=user_id,
        style_profile_id=style.id,
        title=task_input.title,
        genre=task_input.genre,
        content="\n\n".join(paragraphs),
        status="completed",
    )
    document.paragraphs = [
        models.DocumentParagraph(id=new_id("dpara"), position=index, content=paragraph)
        for index, paragraph in enumerate(paragraphs, start=1)
    ]
    writing_task = models.WritingTask(
        id=new_id("task"),
        user_id=user_id,
        style_profile_id=style.id,
        document_id=document.id,
        status="completed",
        genre=task_input.genre,
        title=task_input.title,
        brief=task_input.brief,
        effective_mode=prompt.mode.value,
        rag_enabled=prompt.rag_enabled,
        prompt_version=prompt.prompt_version,
        requirements={
            "genre": task_input.genre,
            "task_type": task_input.task_type,
            "title": task_input.title,
            "brief": task_input.brief,
            "target_length": task_input.target_length,
            "target_reader": task_input.target_reader,
            "must_include": task_input.must_include,
            "must_avoid": task_input.must_avoid,
            "eval_focus": task_input.eval_focus,
            "style_intensity": task_input.style_intensity,
        },
        model_provider=model_result.model_provider,
        model_name=model_result.model_name,
        input_token_count=model_result.input_token_count,
        output_token_count=model_result.output_token_count,
    )
    usage = models.ModelUsageLog(
        id=new_id("usage"),
        user_id=user_id,
        purpose="writing",
        model_provider=model_result.model_provider,
        model_name=model_result.model_name,
        input_token_count=model_result.input_token_count,
        output_token_count=model_result.output_token_count,
    )
    db.add_all([document, writing_task, usage])
    db.commit()
    db.refresh(document)
    db.refresh(writing_task)
    return CreatedWriting(task=writing_task, document=document, model_result=model_result)


def create_free_writing(
    db: Session,
    *,
    user_id: str,
    task_input: WritingTaskInput,
    requested_mode: GenerationMode | str | None = None,
    rag_snippets: list[str] | None = None,
) -> CreatedWriting:
    """自由写作（无风格生成）：不绑定任何用户风格档案。

    style_profile_id 指向系统占位风格（SYSTEM_FREE_WRITE_STYLE_ID），以满足
    documents.style_profile_id NOT NULL 约束；鉴评逻辑据此跳过无风格文章。
    """
    prompt = compose_free_prompt(task=task_input)
    fallback = _fallback_article(style_name="自由写作", task=task_input)
    model_result = ModelGateway().generate(
        messages=[
            {"role": "system", "content": prompt.system_prompt},
            {"role": "user", "content": prompt.user_prompt},
        ],
        purpose="writing",
        fallback=fallback,
    )
    paragraphs = split_paragraphs(model_result.content)
    document = models.Document(
        id=new_id("doc"),
        user_id=user_id,
        style_profile_id=SYSTEM_FREE_WRITE_STYLE_ID,
        title=task_input.title,
        genre=task_input.genre,
        content="\n\n".join(paragraphs),
        status="completed",
    )
    document.paragraphs = [
        models.DocumentParagraph(id=new_id("dpara"), position=index, content=paragraph)
        for index, paragraph in enumerate(paragraphs, start=1)
    ]
    writing_task = models.WritingTask(
        id=new_id("task"),
        user_id=user_id,
        style_profile_id=SYSTEM_FREE_WRITE_STYLE_ID,
        document_id=document.id,
        status="completed",
        genre=task_input.genre,
        title=task_input.title,
        brief=task_input.brief,
        effective_mode=prompt.mode.value,
        rag_enabled=prompt.rag_enabled,
        prompt_version=prompt.prompt_version,
        requirements={
            "genre": task_input.genre,
            "task_type": task_input.task_type,
            "title": task_input.title,
            "brief": task_input.brief,
            "target_length": task_input.target_length,
            "target_reader": task_input.target_reader,
            "must_include": task_input.must_include,
            "must_avoid": task_input.must_avoid,
            "eval_focus": task_input.eval_focus,
            "style_intensity": task_input.style_intensity,
        },
        model_provider=model_result.model_provider,
        model_name=model_result.model_name,
        input_token_count=model_result.input_token_count,
        output_token_count=model_result.output_token_count,
    )
    usage = models.ModelUsageLog(
        id=new_id("usage"),
        user_id=user_id,
        purpose="writing",
        model_provider=model_result.model_provider,
        model_name=model_result.model_name,
        input_token_count=model_result.input_token_count,
        output_token_count=model_result.output_token_count,
    )
    db.add_all([document, writing_task, usage])
    db.commit()
    db.refresh(document)
    db.refresh(writing_task)
    return CreatedWriting(task=writing_task, document=document, model_result=model_result)


def preview_rewrite_paragraph(
    db: Session,
    *,
    user_id: str,
    document_id: str,
    paragraph_id: str,
    instruction: str,
) -> dict[str, Any]:
    """Call AI to rewrite a paragraph and return the result WITHOUT saving to DB.

    Returns a dict with the rewritten content plus token counts for cost accounting.
    """
    document = db.scalar(
        select(models.Document)
        .where(models.Document.id == document_id, models.Document.user_id == user_id)
        .options(selectinload(models.Document.paragraphs))
    )
    if not document:
        raise HTTPException(status_code=404, detail="document not found")
    paragraph = next((item for item in document.paragraphs if item.id == paragraph_id), None)
    if not paragraph:
        raise HTTPException(status_code=404, detail="paragraph not found")
    style = get_active_style(db, user_id=user_id, style_profile_id=document.style_profile_id)
    fallback = f"{paragraph.content}\n（已按意见调整：{instruction.strip()}；保持\u201c{style.name}\u201d的语气和自然段长度。）"
    model_result = ModelGateway().generate(
        messages=[
            {"role": "system", "content": "你只重写用户指定的一个自然段，其他段落不得改变。"},
            {
                "role": "user",
                "content": f"风格：{style.profile}\n原段落：{paragraph.content}\n修改意见：{instruction}",
            },
        ],
        purpose="paragraph_rewrite",
        fallback=fallback,
    )
    db.add(
        models.ModelUsageLog(
            id=new_id("usage"),
            user_id=user_id,
            purpose="paragraph_rewrite",
            model_provider=model_result.model_provider,
            model_name=model_result.model_name,
            input_token_count=model_result.input_token_count,
            output_token_count=model_result.output_token_count,
        )
    )
    db.commit()
    return {
        "content": model_result.content.strip(),
        "input_token_count": model_result.input_token_count,
        "output_token_count": model_result.output_token_count,
        "model_name": model_result.model_name,
    }


def update_paragraph_content(
    db: Session,
    *,
    user_id: str,
    document_id: str,
    paragraph_id: str,
    content: str,
) -> models.Document:
    """Directly update a paragraph's content (manual edit or confirmed AI rewrite)."""
    document = db.scalar(
        select(models.Document)
        .where(models.Document.id == document_id, models.Document.user_id == user_id)
        .options(selectinload(models.Document.paragraphs))
    )
    if not document:
        raise HTTPException(status_code=404, detail="document not found")
    paragraph = next((item for item in document.paragraphs if item.id == paragraph_id), None)
    if not paragraph:
        raise HTTPException(status_code=404, detail="paragraph not found")
    paragraph.content = content.strip()
    paragraph.rewrite_count += 1
    paragraph.updated_at = models.utc_now()
    ordered = sorted(document.paragraphs, key=lambda item: item.position)
    document.content = "\n\n".join(item.content for item in ordered)
    document.updated_at = models.utc_now()
    db.commit()
    db.refresh(document)
    return document


def revise_document(
    db: Session,
    *,
    user_id: str,
    document_id: str,
    instruction: str,
) -> tuple[models.Document, ModelResult]:
    """无风格自由写作文章的「继续修改」：用当前文章 + 原始写作要求 + 新修改意见重生成，覆盖当前文档。

    旧段落会被删除（关系级 delete-orphan）并以新段落覆盖；新建一条 task_type=改写 的
    WritingTask 用于审计。自由写作不鉴评。返回更新后的文档与本次模型结果（供积分扣减）。
    """
    document = get_user_document(db, user_id=user_id, document_id=document_id)
    task = db.scalar(
        select(models.WritingTask)
        .where(models.WritingTask.document_id == document.id, models.WritingTask.user_id == user_id)
        .order_by(models.WritingTask.created_at.desc())
    )
    requirements = task.requirements if task else {}
    genre = document.genre
    title = document.title
    target_length = requirements.get("target_length") or f"约 {len(document.content)} 字"
    original_brief = (task.brief if task else "") or ""

    system_prompt = (
        "你是个人风格写作 Agent。你的任务是根据用户给出的修改意见，重写整篇文章。"
        "严格遵循原文的写作要求（文体、标题、大体长度），只依据修改意见调整，不额外发挥、不解释过程。"
        "直接输出新的文章正文，用空行分隔自然段，不要任何前缀、标题或说明文字。"
    )
    user_prompt = (
        "## 原始写作要求\n"
        f"- 文体：{genre}\n"
        f"- 标题/主题：{title}\n"
        f"- 目标长度：{target_length}\n"
        f"- 原始需求：{original_brief}\n\n"
        "## 当前文章\n"
        f"{document.content}\n\n"
        "## 用户的修改意见\n"
        f"{instruction}\n\n"
        "请根据以上修改意见重新生成完整文章，保持相同文体与大体长度，直接输出正文。"
    )
    fallback = document.content
    model_result = ModelGateway().generate(
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        purpose="writing",
        fallback=fallback,
    )
    paragraphs = split_paragraphs(model_result.content)
    if not paragraphs:
        paragraphs = split_paragraphs(document.content)

    # 覆盖旧段落（relationship 的 delete-orphan 会清理原段落）
    document.paragraphs = [
        models.DocumentParagraph(id=new_id("dpara"), position=index, content=paragraph)
        for index, paragraph in enumerate(paragraphs, start=1)
    ]
    document.content = "\n\n".join(paragraphs)
    document.updated_at = models.utc_now()

    writing_task = models.WritingTask(
        id=new_id("task"),
        user_id=user_id,
        style_profile_id=document.style_profile_id,
        document_id=document.id,
        status="completed",
        genre=genre,
        title=title,
        brief=f"改写：{instruction}",
        effective_mode="revise",
        rag_enabled=False,
        prompt_version="revise_v1",
        requirements={
            "genre": genre,
            "task_type": "改写",
            "title": title,
            "target_length": target_length,
            "original_brief": original_brief,
            "instruction": instruction,
        },
        model_provider=model_result.model_provider,
        model_name=model_result.model_name,
        input_token_count=model_result.input_token_count,
        output_token_count=model_result.output_token_count,
    )
    usage = models.ModelUsageLog(
        id=new_id("usage"),
        user_id=user_id,
        purpose="writing",
        model_provider=model_result.model_provider,
        model_name=model_result.model_name,
        input_token_count=model_result.input_token_count,
        output_token_count=model_result.output_token_count,
    )
    db.add_all([writing_task, usage])
    db.commit()
    db.refresh(document)
    return document, model_result


def save_document(db: Session, *, user_id: str, document_id: str) -> models.Document:
    document = db.scalar(
        select(models.Document)
        .where(models.Document.id == document_id, models.Document.user_id == user_id)
        .options(selectinload(models.Document.paragraphs))
    )
    if not document:
        raise HTTPException(status_code=404, detail="document not found")
    document.is_saved = True
    document.saved_at = models.utc_now()
    db.commit()
    db.refresh(document)
    return document


def unsave_document(db: Session, *, user_id: str, document_id: str) -> models.Document:
    document = db.scalar(
        select(models.Document)
        .where(models.Document.id == document_id, models.Document.user_id == user_id)
        .options(selectinload(models.Document.paragraphs))
    )
    if not document:
        raise HTTPException(status_code=404, detail="document not found")
    document.is_saved = False
    document.saved_at = None
    db.commit()
    db.refresh(document)
    return document


def get_user_document(db: Session, *, user_id: str, document_id: str) -> models.Document:
    document = db.scalar(
        select(models.Document)
        .where(models.Document.id == document_id, models.Document.user_id == user_id)
        .options(selectinload(models.Document.paragraphs))
    )
    if not document:
        raise HTTPException(status_code=404, detail="document not found")
    return document


def list_saved_documents(db: Session, *, user_id: str) -> list[models.Document]:
    return list(
        db.scalars(
            select(models.Document)
            .where(models.Document.user_id == user_id, models.Document.is_saved.is_(True))
            .order_by(models.Document.saved_at.desc())
        )
    )


def generate_docx_bytes(document: models.Document) -> bytes:
    docx = DocxDocument()
    title = docx.add_heading(document.title, level=1)
    title.alignment = 1  # center
    meta = docx.add_paragraph()
    meta.alignment = 1
    run = meta.add_run(f"{document.genre} · 约 {len(document.content)} 字 · {document.updated_at.strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    docx.add_paragraph()
    for paragraph in sorted(document.paragraphs, key=lambda item: item.position):
        p = docx.add_paragraph(paragraph.content)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(8)
    import io

    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _fallback_article(*, style_name: str, task: WritingTaskInput) -> str:
    target = task.target_length.strip() or "1200字"
    requested_digits = "".join(ch for ch in target if ch.isdigit())
    requested_length = int(requested_digits) if requested_digits else 600
    paragraph_count = 4 if requested_length < 1000 else 6
    variant_index = next(FALLBACK_GENERATION_COUNTER)
    variant_markers = ["木尺", "搪瓷杯", "旧雨伞", "蓝布围裙", "纸绳", "玻璃糖罐"]
    variant_marker = variant_markers[variant_index % len(variant_markers)]
    if "诗" in task.genre or "行" in target:
        line_sets = [
            [
                f"{task.title}落在傍晚的门边，",
                "灯光低下来，像有人把话说轻。",
                f"{variant_marker}没有急着证明自己，",
                "只在木纹和灰尘里留住一小段时间。",
                "街角的声音经过，又慢慢散开，",
                "我站了一会儿，没有把它叫作怀念。",
            ],
            [
                "雨停以后，门槛还湿着，",
                f"{variant_marker}靠在墙边，颜色暗下去。",
                "有人从小店里出来，",
                "把找回的零钱握得很轻。",
                "那一点声响没有留下来，",
                "只让傍晚慢了一步。",
            ],
        ]
        return "\n".join(line_sets[variant_index % len(line_sets)])

    include_terms = [item.strip() for item in task.must_include.replace("，", ",").split(",") if item.strip()]
    scene_hint = include_terms[0] if include_terms else "街角"
    openings = [
        f"早晨过去一半时，{task.title}还没有完全醒来。{scene_hint}停在路边，卷帘门拉到一半，里面有纸箱、旧凳子和一点隔夜的灰。有人推门进去，又很快出来，手里多了一袋东西，脚步没有惊动门口晒太阳的猫。",
        f"雨停以后，{task.title}显得比平时低一些。{scene_hint}贴着墙根展开，门口积着浅浅的水，{variant_marker}靠在柜台旁，像刚被人放下，又像已经在那里等了很久。",
        f"傍晚还没落下来，{task.title}先有了旧物的气味。{scene_hint}不在远处，就在门槛、纸箱和半开的抽屉之间。有人弯腰找东西，衣角碰到木凳，声音轻得像没有发生。",
    ]

    base_paragraphs = [
        openings[variant_index % len(openings)],
        f"早晨过去一半时，{task.title}还没有完全醒来。{scene_hint}停在路边，卷帘门拉到一半，里面有纸箱、旧凳子和一点隔夜的灰。有人推门进去，又很快出来，手里多了一袋东西，脚步没有惊动门口晒太阳的猫。",
        f"我常从那里经过。玻璃柜台擦得不算亮，角落里压着几张发黄的收据，墙上的钟慢了几分钟。店主说话不急，找零钱时先把纸币抹平，再把硬币放到掌心。那些动作很小，却让人觉得日子仍有可以安放的地方。",
        f"旧物在这里不是摆设。搪瓷杯缺了一点边，竹椅的扶手被磨出暗光，门后的雨伞有两根伞骨微微弯着。它们不解释自己的来处，只在被拿起、放下、再次靠回墙边的时候，露出一点被时间用过的痕迹。",
        f"到了午后，街上的车声变密，店里反而安静。有人来买针线，有人问一节电池，也有人只是站在门口避一阵太阳。店主没有多问，把风扇转过去一点。那一刻，空气里有尘土、塑料袋和旧木头混在一起的味道，并不新鲜，却具体。",
        f"我离开时，门口的小灯还没有亮。身后传来卷尺收回去的声音，很轻，像一句没有说完的话。这样的地方不会替谁总结生活，只把零散的需要接住一点，再让人继续往前走。",
        f"后来我想起它，想起的也不是完整故事，只是柜台边那束斜光，和店主把零钱递出来时停顿的一瞬。附近生活大概就是这样，不声张，也不空着，在人们低头寻找一件小东西的时候，慢慢显出自己的重量。",
    ]
    selected = base_paragraphs[:paragraph_count]
    return "\n\n".join(selected)





