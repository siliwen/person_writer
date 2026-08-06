from __future__ import annotations

import argparse
import io
import ipaddress
import json
import re
import socket
import sys
from dataclasses import asdict, dataclass
from datetime import datetime
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.parse import urljoin, urlparse
from urllib.request import HTTPRedirectHandler, Request, build_opener

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree, html


APP_DIR = Path(__file__).resolve().parent
STATIC_DIR = APP_DIR / "static"
PROJECT_DIR = APP_DIR.parents[1]
EVAL_OUTPUT_DIR = PROJECT_DIR / "测评集" / "网页作品"
MAX_RESPONSE_BYTES = 5 * 1024 * 1024
MAX_REDIRECTS = 5
FETCH_TIMEOUT_SECONDS = 15
USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36 "
    "ArticleCollector/1.0"
)

SPACE_RE = re.compile(r"[\t\r\f\v ]+")
BLANK_RE = re.compile(r"\n{3,}")
PUNCTUATION_RE = re.compile(r"[，。！？；：,.!?;:]")
POSITIVE_HINT_RE = re.compile(
    r"article|content|entry|main|post|story|正文|文章|详情|新闻", re.I
)
NEGATIVE_HINT_RE = re.compile(
    r"comment|footer|header|menu|nav|related|recommend|share|sidebar|"
    r"social|subscribe|广告|推荐|评论|导航|页脚|侧栏|分享",
    re.I,
)
INVALID_FILENAME_RE = re.compile(r'[\\/:*?"<>|\x00-\x1f]')
DATE_TEXT_RE = re.compile(r"\d{4}年?\d{1,2}月?\d{1,2}日?(?:\s*\d{1,2}:\d{2})?")


class UserFacingError(Exception):
    pass


@dataclass
class ContentBlock:
    type: str
    text: str


@dataclass
class Article:
    title: str
    author: str
    published_at: str
    source_url: str
    site_name: str
    blocks: list[ContentBlock]

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks)

    @property
    def char_count(self) -> int:
        return len(re.sub(r"\s", "", self.text))

    def to_dict(self) -> dict[str, Any]:
        result = asdict(self)
        result["text"] = self.text
        result["char_count"] = self.char_count
        return result


def normalize_text(value: str | None) -> str:
    if not value:
        return ""
    lines = [SPACE_RE.sub(" ", line).strip() for line in value.splitlines()]
    return BLANK_RE.sub("\n\n", "\n".join(line for line in lines if line)).strip()


def is_public_ip(address: str) -> bool:
    ip = ipaddress.ip_address(address)
    return ip.is_global


def validate_public_url(raw_url: str) -> str:
    raw_url = raw_url.strip()
    if not raw_url:
        raise UserFacingError("请输入网页地址。")
    if "://" not in raw_url:
        raw_url = "https://" + raw_url

    parsed = urlparse(raw_url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise UserFacingError("仅支持有效的 http 或 https 网页地址。")
    if parsed.username or parsed.password:
        raise UserFacingError("网址中不能包含用户名或密码。")

    try:
        addresses = {
            item[4][0]
            for item in socket.getaddrinfo(parsed.hostname, parsed.port or 443)
        }
    except socket.gaierror as exc:
        raise UserFacingError("无法解析该网址的域名。") from exc

    if not addresses or any(not is_public_ip(address) for address in addresses):
        raise UserFacingError("出于安全考虑，不能抓取本机或内网地址。")
    return parsed.geturl()


class SafeRedirectHandler(HTTPRedirectHandler):
    def __init__(self) -> None:
        self.redirect_count = 0
        super().__init__()

    def redirect_request(self, req, fp, code, msg, headers, newurl):
        self.redirect_count += 1
        if self.redirect_count > MAX_REDIRECTS:
            raise UserFacingError("网页跳转次数过多，已停止抓取。")
        safe_url = validate_public_url(urljoin(req.full_url, newurl))
        return super().redirect_request(req, fp, code, msg, headers, safe_url)


def decode_html(payload: bytes, content_type: str) -> str:
    charset_match = re.search(r"charset\s*=\s*[\"']?([^\s;\"']+)", content_type, re.I)
    candidates = [charset_match.group(1)] if charset_match else []
    head = payload[:4096].decode("ascii", errors="ignore")
    meta_match = re.search(r"charset\s*=\s*[\"']?([^\s/>\"']+)", head, re.I)
    if meta_match:
        candidates.append(meta_match.group(1))
    candidates.extend(["utf-8", "gb18030"])

    for encoding in dict.fromkeys(candidates):
        try:
            return payload.decode(encoding)
        except (LookupError, UnicodeDecodeError):
            continue
    return payload.decode("utf-8", errors="replace")


def fetch_html(raw_url: str) -> tuple[str, str]:
    url = validate_public_url(raw_url)
    request = Request(
        url,
        headers={
            "User-Agent": USER_AGENT,
            "Accept": "text/html,application/xhtml+xml;q=0.9,*/*;q=0.1",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.7",
        },
    )
    opener = build_opener(SafeRedirectHandler())

    try:
        with opener.open(request, timeout=FETCH_TIMEOUT_SECONDS) as response:
            content_type = response.headers.get("Content-Type", "")
            if "text/html" not in content_type and "application/xhtml+xml" not in content_type:
                raise UserFacingError("该网址返回的不是 HTML 网页。")
            content_length = response.headers.get("Content-Length")
            if content_length and int(content_length) > MAX_RESPONSE_BYTES:
                raise UserFacingError("网页内容超过 5 MB，已停止抓取。")
            payload = response.read(MAX_RESPONSE_BYTES + 1)
            if len(payload) > MAX_RESPONSE_BYTES:
                raise UserFacingError("网页内容超过 5 MB，已停止抓取。")
            final_url = validate_public_url(response.geturl())
            return decode_html(payload, content_type), final_url
    except UserFacingError:
        raise
    except HTTPError as exc:
        raise UserFacingError(f"网页返回 HTTP {exc.code}，无法抓取。") from exc
    except URLError as exc:
        raise UserFacingError("网页连接失败，请检查网址或稍后重试。") from exc
    except TimeoutError as exc:
        raise UserFacingError("网页响应超时，请稍后重试。") from exc


def first_meta(tree: html.HtmlElement, *selectors: str) -> str:
    for selector in selectors:
        values = tree.xpath(selector)
        for value in values:
            text = normalize_text(str(value))
            if text:
                return text
    return ""


def extract_title(tree: html.HtmlElement) -> str:
    title = first_meta(
        tree,
        "//*[@id='newstit']//text()",
        "//meta[@property='og:title']/@content",
        "//meta[@name='twitter:title']/@content",
        "//main//*[self::h1 or self::h2 or self::h3 or self::h4 or self::h5 or self::h6][1]//text()",
    )
    if title:
        return title
    browser_title = first_meta(tree, "//title/text()")
    return re.split(r"\s*(?:--|—|_\s*)\s*", browser_title, maxsplit=1)[0].strip()


def extract_author(tree: html.HtmlElement) -> str:
    author_candidates = [
        first_meta(tree, "//meta[@name='author']/@content"),
        first_meta(tree, "//meta[@property='article:author']/@content"),
        first_meta(tree, "//*[@rel='author'][1]//text()"),
        first_meta(tree, "//*[contains(@class, 'author')][1]//text()"),
    ]
    for author in author_candidates:
        if author and not author.isdigit() and len(author) <= 80:
            return author

    info_lines = tree.xpath(
        "//*[contains(concat(' ', normalize-space(@class), ' '), ' end_info ')]"
    )
    for node in info_lines:
        text = normalize_text(node.text_content())
        parts = [normalize_text(part) for part in re.split(r"[|｜]", text)]
        if len(parts) >= 2:
            candidate = DATE_TEXT_RE.sub("", parts[1]).strip(" -·　")
            if candidate and not candidate.isdigit() and len(candidate) <= 80:
                return candidate

    body_text = normalize_text(tree.text_content())
    match = re.search(r"(?:作者|撰文)[：:]\s*([^\s|｜]{2,40})", body_text)
    return match.group(1).strip() if match else ""


def remove_noise(tree: html.HtmlElement) -> None:
    selectors = [
        "//script", "//style", "//noscript", "//template", "//svg", "//canvas",
        "//form", "//nav", "//footer", "//aside", "//iframe", "//button",
        "//*[contains(translate(@aria-hidden, 'TRUE', 'true'), 'true')]",
    ]
    for node in tree.xpath(" | ".join(selectors)):
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)


def link_density(node: html.HtmlElement) -> float:
    text_length = len(normalize_text(node.text_content()))
    if not text_length:
        return 1.0
    link_length = sum(len(normalize_text(link.text_content())) for link in node.xpath(".//a"))
    return min(1.0, link_length / text_length)


def candidate_score(node: html.HtmlElement) -> float:
    text = normalize_text(node.text_content())
    if len(text) < 120:
        return float("-inf")
    paragraphs = [normalize_text(p.text_content()) for p in node.xpath(".//p")]
    useful_paragraphs = [p for p in paragraphs if len(p) >= 20]
    hints = " ".join([node.get("id", ""), node.get("class", ""), node.tag])
    score = min(len(text), 12000) / 35
    score += len(useful_paragraphs) * 28
    score += len(PUNCTUATION_RE.findall(text)) * 1.4
    score -= link_density(node) * 420
    if node.tag in {"article", "main"}:
        score += 260
    if POSITIVE_HINT_RE.search(hints):
        score += 160
    if NEGATIVE_HINT_RE.search(hints):
        score -= 280
    return score


def select_main_node(tree: html.HtmlElement) -> html.HtmlElement:
    preferred = tree.xpath(
        "//*[contains(concat(' ', normalize-space(@class), ' '), ' end_article ')]"
        " | //article | //main | //*[@role='main']"
    )
    candidates = preferred + tree.xpath("//section | //div")
    scored = [(candidate_score(node), node) for node in candidates]
    scored = [item for item in scored if item[0] != float("-inf")]
    if scored:
        return max(scored, key=lambda item: item[0])[1]
    body = tree.find("body")
    return body if body is not None else tree


def extract_blocks(main_node: html.HtmlElement, title: str) -> list[ContentBlock]:
    blocks: list[ContentBlock] = []
    seen: set[str] = set()
    for node in main_node.xpath(".//h1 | .//h2 | .//h3 | .//p | .//blockquote | .//li"):
        if any(parent.tag in {"nav", "footer", "aside"} for parent in node.iterancestors()):
            continue
        text = normalize_text(node.text_content())
        key = re.sub(r"\s", "", text)
        if not text or len(key) < 2 or key in seen:
            continue
        if title and key == re.sub(r"\s", "", title):
            continue
        if node.tag in {"p", "li"} and len(key) < 8:
            continue
        seen.add(key)
        block_type = "heading" if node.tag in {"h1", "h2", "h3"} else "paragraph"
        blocks.append(ContentBlock(block_type, text))

    if not blocks:
        fallback = normalize_text(main_node.text_content())
        blocks = [ContentBlock("paragraph", part) for part in fallback.split("\n\n") if part]
    return blocks


def parse_article(markup: str, source_url: str) -> Article:
    parser = html.HTMLParser(encoding="utf-8", recover=True)
    try:
        tree = html.fromstring(markup.encode("utf-8"), parser=parser, base_url=source_url)
    except (etree.ParserError, ValueError) as exc:
        raise UserFacingError("网页 HTML 无法解析。") from exc

    title = extract_title(tree)
    author = extract_author(tree)
    published_at = first_meta(
        tree,
        "//meta[@property='article:published_time']/@content",
        "//meta[@name='publishdate']/@content",
        "//meta[@name='date']/@content",
        "//time[1]/@datetime",
        "//time[1]//text()",
    )
    site_name = first_meta(tree, "//meta[@property='og:site_name']/@content")

    remove_noise(tree)
    main_node = select_main_node(tree)
    blocks = extract_blocks(main_node, title)
    char_count = len(re.sub(r"\s", "", "".join(block.text for block in blocks)))
    if char_count < 80:
        raise UserFacingError(
            "没有识别到足够的正文。该网页可能需要登录，或正文由 JavaScript 动态加载。"
        )
    return Article(
        title=title or "未命名网页作品",
        author=author,
        published_at=published_at,
        source_url=source_url,
        site_name=site_name,
        blocks=blocks,
    )


def article_name_without_author(article: Article) -> str:
    title = article.title.strip()
    if article.author:
        pattern = rf"^{re.escape(article.author.strip())}\s*[：:]\s*"
        title = re.sub(pattern, "", title, count=1)
    return title or article.title.strip() or "未命名网页作品"


def safe_filename_part(value: str, max_length: int = 80) -> str:
    cleaned = INVALID_FILENAME_RE.sub("_", normalize_text(value))
    cleaned = cleaned.rstrip(". ").strip()
    return cleaned[:max_length].rstrip(". ") or "未知"


def filename_for_article(article: Article) -> str:
    author = safe_filename_part(article.author or "未知作者", 40)
    article_name = safe_filename_part(article_name_without_author(article), 100)
    return f"{author}+{article_name}.docx"


def save_article_docx(
    article: Article, output_dir: Path = EVAL_OUTPUT_DIR
) -> tuple[Path, bytes]:
    body = article_to_docx(article)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename_for_article(article)
    temporary_path = output_path.with_suffix(".docx.tmp")
    temporary_path.write_bytes(body)
    temporary_path.replace(output_path)
    return output_path, body


def set_run_font(run, ascii_font: str, east_asia_font: str, size: float, *, bold=False, color=None):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2F6F64")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def article_to_docx(article: Article) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33

    heading = document.styles["Heading 2"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(31, 77, 70)
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(6)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("网页作品采集")
    set_run_font(run, "Calibri", "Microsoft YaHei", 9.5, bold=True, color="2F6F64")

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(10)
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(article.title)
    set_run_font(title_run, "Calibri", "Microsoft YaHei", 24, bold=True, color="17201E")

    meta_parts = [part for part in [article.author, article.published_at, article.site_name] if part]
    if meta_parts:
        meta = document.add_paragraph()
        meta.paragraph_format.space_after = Pt(4)
        meta_run = meta.add_run("  |  ".join(meta_parts))
        set_run_font(meta_run, "Calibri", "Microsoft YaHei", 9.5, color="5F6B67")

    source = document.add_paragraph()
    source.paragraph_format.space_after = Pt(20)
    label = source.add_run("来源：")
    set_run_font(label, "Calibri", "Microsoft YaHei", 9, color="5F6B67")
    add_hyperlink(source, article.source_url, article.source_url)

    for block in article.blocks:
        if block.type == "heading":
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.add_run(block.text)
        else:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.widow_control = True
            paragraph.add_run(block.text)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"采集于 {datetime.now().strftime('%Y-%m-%d')}")
    set_run_font(footer_run, "Calibri", "Microsoft YaHei", 8, color="7A8581")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def article_from_payload(payload: dict[str, Any]) -> Article:
    raw_blocks = payload.get("blocks")
    if not isinstance(raw_blocks, list) or not raw_blocks:
        raise UserFacingError("没有可导出的正文内容。")
    blocks = []
    for raw_block in raw_blocks:
        if not isinstance(raw_block, dict):
            continue
        text = normalize_text(str(raw_block.get("text", "")))
        block_type = "heading" if raw_block.get("type") == "heading" else "paragraph"
        if text:
            blocks.append(ContentBlock(block_type, text))
    if not blocks:
        raise UserFacingError("没有可导出的正文内容。")
    return Article(
        title=normalize_text(str(payload.get("title", ""))) or "未命名网页作品",
        author=normalize_text(str(payload.get("author", ""))),
        published_at=normalize_text(str(payload.get("published_at", ""))),
        source_url=validate_public_url(str(payload.get("source_url", ""))),
        site_name=normalize_text(str(payload.get("site_name", ""))),
        blocks=blocks,
    )


class AppHandler(SimpleHTTPRequestHandler):
    server_version = "ArticleCollector/1.0"

    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(STATIC_DIR), **kwargs)

    def log_message(self, format: str, *args) -> None:
        sys.stdout.write("[%s] %s\n" % (self.log_date_time_string(), format % args))

    def send_json(self, payload: dict[str, Any], status: int = 200) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Cache-Control", "no-store")
        self.end_headers()
        self.wfile.write(body)

    def read_json(self) -> dict[str, Any]:
        content_length = int(self.headers.get("Content-Length", "0"))
        if content_length <= 0 or content_length > MAX_RESPONSE_BYTES:
            raise UserFacingError("请求内容为空或过大。")
        try:
            payload = json.loads(self.rfile.read(content_length).decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise UserFacingError("请求数据格式不正确。") from exc
        if not isinstance(payload, dict):
            raise UserFacingError("请求数据格式不正确。")
        return payload

    def do_GET(self) -> None:
        if self.path == "/api/health":
            self.send_json({"status": "ok"})
            return
        if self.path == "/":
            self.path = "/index.html"
        super().do_GET()

    def do_POST(self) -> None:
        try:
            if self.path == "/api/extract":
                payload = self.read_json()
                markup, final_url = fetch_html(str(payload.get("url", "")))
                article = parse_article(markup, final_url)
                self.send_json({"ok": True, "article": article.to_dict()})
                return
            if self.path == "/api/export":
                article = article_from_payload(self.read_json())
                output_path, _ = save_article_docx(article)
                self.send_json(
                    {
                        "ok": True,
                        "saved_filename": output_path.name,
                        "saved_directory": "测评集/网页作品",
                    }
                )
                return
            self.send_json({"ok": False, "error": "接口不存在。"}, HTTPStatus.NOT_FOUND)
        except UserFacingError as exc:
            self.send_json({"ok": False, "error": str(exc)}, HTTPStatus.BAD_REQUEST)
        except Exception as exc:
            print(f"Unexpected error: {exc!r}", file=sys.stderr)
            self.send_json(
                {"ok": False, "error": "处理失败，请查看服务端日志。"},
                HTTPStatus.INTERNAL_SERVER_ERROR,
            )


def main() -> None:
    parser = argparse.ArgumentParser(description="网页正文采集与 Word 导出工具")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8765)
    args = parser.parse_args()
    server = ThreadingHTTPServer((args.host, args.port), AppHandler)
    print(f"网页正文采集工具已启动：http://{args.host}:{args.port}")
    print("按 Ctrl+C 停止服务。")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()


if __name__ == "__main__":
    main()
